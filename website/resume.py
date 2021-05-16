"""Generating a DOCX and converting it to .pdf through python"""
from .resources import get_resource_json
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn as sharedqn
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches
from os.path import expanduser, join
from sys import platform

if platform == "linux":
    RESUME_DIRECTORY_LOCATION = "/tmp/"
else:
    RESUME_DIRECTORY_LOCATION = expanduser('~')

RESUME_FILENAME = 'Derek Santos - Resume.docx'
RESUME_LOCATION = join(RESUME_DIRECTORY_LOCATION, RESUME_FILENAME)
DEFAULT_SPACING = Cm(0.03)
DEFAULT_FONT_NAME = "Calibri Light"
DEFAULT_FONT_COLOR = RGBColor(0, 0, 0)
DEFAULT_FONT_SIZE_TITLE = Pt(18)
DEFAULT_FONT_SIZE_SUBTITLE = Pt(8)
DEFAULT_FONT_SIZE_HEADING = Pt(8.8)
DEFAULT_FONT_SIZE_TEXT = Pt(7.5)
DEFAULT_TOP_MARGIN_LENGTH = Cm(1)
DEFAULT_BOTTOM_MARGIN_LENGTH = Cm(1)
DEFAULT_LEFT_MARGIN_LENGTH = Inches(.75)
DEFAULT_RIGHT_MARGIN_LENGTH = Inches(.75)


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(sharedqn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = OxmlElement('w:color')
        c.set(sharedqn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(sharedqn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def style_document(document):
    # Style -> Title
    style = document.styles.add_style('ResumeTitle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_TITLE
    font.color.rgb = DEFAULT_FONT_COLOR
    formatting = style.paragraph_format
    formatting.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formatting.space_before = DEFAULT_SPACING
    formatting.space_after = DEFAULT_SPACING

    # Style -> Subtitle
    style = document.styles.add_style('ResumeSubtitle',
                                      WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_SUBTITLE
    font.color.rgb = RGBColor(128, 128, 128)
    formatting = style.paragraph_format
    formatting.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formatting.space_before = DEFAULT_SPACING
    formatting.space_after = DEFAULT_SPACING

    # Style -> Normal
    style = document.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_TEXT
    font.color.rgb = DEFAULT_FONT_COLOR
    formatting = style.paragraph_format
    formatting.space_before = DEFAULT_SPACING
    formatting.space_after = DEFAULT_SPACING

    # Style -> Header
    style = document.styles.add_style(
        'ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_HEADING
    font.bold = True
    font.color.rgb = DEFAULT_FONT_COLOR
    formatting = style.paragraph_format
    formatting.space_before = DEFAULT_SPACING
    formatting.space_after = DEFAULT_SPACING

    # Style -> List Bullet
    style = document.styles['List Bullet']
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_TEXT
    font.color.rgb = DEFAULT_FONT_COLOR
    formatting = style.paragraph_format
    formatting.space_before = DEFAULT_SPACING
    formatting.space_after = DEFAULT_SPACING
    formatting.left_indent = Inches(0.5)

    # Style -> Emphasis
    style = document.styles['Emphasis']
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE_TEXT
    font.color.rgb = RGBColor(84, 84, 84)


def technical_skills(document):
    head = document.add_paragraph('Technical Skills',
                                  style='ResumeHeader')
    insertHR(head)
    # Technologies
    document.add_paragraph('').add_run('Technologies').bold = True
    technologies = document.add_paragraph('')
    skills_json = get_resource_json('skills.json')
    technology_list = skills_json['technologies']
    technology_list.sort(key=lambda x: x[1], reverse=True)
    for pos, item in enumerate(technology_list):
        if pos != 0:
            technologies.add_run(', ')
        technologies.add_run('{}'.format(item[0]))

    # Tools
    document.add_paragraph('').add_run('Tools').bold = True
    tools = document.add_paragraph('')
    tools_list = skills_json['tools']
    tools_list.sort(key=lambda x: x[1], reverse=True)
    for pos, item in enumerate(tools_list):
        if pos != 0:
            tools.add_run(', ')
        tools.add_run('{}'.format(item[0]))


def experience(document):
    head = document.add_paragraph('Experience',
                                  style='ResumeHeader')
    insertHR(head)
    experiences = get_resource_json("career.json")
    for experience in experiences:
        experience_paragraph = document.add_paragraph('')
        experience_paragraph.add_run(' {}  '.format(experience['date']),
                                     'Emphasis').bold = True
        experience_paragraph.add_run(
            '{}  '.format(experience['title'])).bold = True

        for pos, tech in enumerate(experience['technologies']):
            if pos != 0:
                experience_paragraph.add_run(',', 'Emphasis')
            subtext = experience_paragraph.add_run(' {}'.format(tech),
                                                   'Emphasis')

        experience_paragraph.add_run(' - {}'.format(experience['location']),
                                     'Emphasis')

        # experience bullet points - Has to be on its own paragraph
        for bullet in experience['descriptions']:
            document.add_paragraph(bullet, style='List Bullet')


def leadership(document):
    # Leadership
    head = document.add_paragraph('Leadership',
                                  style='ResumeHeader')
    insertHR(head)
    for item in get_resource_json('leadership.json'):
        leadership_paragraph = document.add_paragraph('')
        date = leadership_paragraph.add_run('{}  '.format(item['date']),
                                            'Emphasis')
        leadership_paragraph.add_run(item['title']).bold = True
        leadership_paragraph.add_run(' {} '.format(item['location']),
                                     'Emphasis')

        date.italic = True
        date.bold = True
        for line in item['description']:
            document.add_paragraph(line, style='List Bullet')


def education(document):
    # Education
    head = document.add_paragraph('Education',
                                  style='ResumeHeader')
    insertHR(head)
    for item in get_resource_json('education.json'):

        title_line = document.add_paragraph('')
        date = title_line.add_run('{}  '.format(item['date']),
                                  'Emphasis')
        title_line.add_run(item['title'] + ' - ').bold = True
        title_line.add_run(item['degree'], "Emphasis")

        date.italic = True
        date.bold = True


def generate_document(location=RESUME_LOCATION):
    """Generate a word document based off the resource documents JSON"""
    # Document Wide Formatting
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = DEFAULT_TOP_MARGIN_LENGTH
        section.bottom_margin = DEFAULT_BOTTOM_MARGIN_LENGTH
        section.left_margin = DEFAULT_LEFT_MARGIN_LENGTH
        section.right_margin = DEFAULT_RIGHT_MARGIN_LENGTH

    style_document(document)

    # Title and sub-heading
    document.add_paragraph('Derek Santos', style='ResumeTitle')

    subtitle = document.add_paragraph('', style='ResumeSubtitle')
    add_hyperlink(subtitle, 'https://santosderek.com/',
                  'santosderek.com',  '568ed2', False)

    subtitle.add_run(' | ')
    subtitle.add_run('Raleigh, NC | santos.jon.derek@gmail.com')

    subtitle.add_run(' | ')
    add_hyperlink(subtitle, 'https://www.linkedin.com/in/santosderek/',
                  'Linkedin',  '568ed2', False)

    subtitle.add_run(' | ')
    add_hyperlink(subtitle, 'https://github.com/santosderek/',
                  'Github',  '568ed2', False)

    technical_skills(document)
    experience(document)
    leadership(document)
    education(document)

    document.save(RESUME_LOCATION)
